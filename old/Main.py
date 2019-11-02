#%% Import all packages in need.
import random
import math
import pandas as pd
import cv2
import os
from shutil import copyfile
from PIL import Image
from translate import Translator
import docx
import ProteinParticle as pp
import Parameters as pa
import Visualization as vis
import Projection as proj
#%% Create lists to store grid positions and remove boundaries.
print('''
      Initializing the run...
''')
listPositions=list(pp.my_grid.grid.keys())
boundaryDeleted=0
for i in range(len(listPositions)):
    if (listPositions[i-boundaryDeleted][0]==0 or listPositions[i-boundaryDeleted][0]==pp.my_grid.bx or listPositions[i-boundaryDeleted][1]==0 or listPositions[i-boundaryDeleted][1]==pp.my_grid.by or listPositions[i-boundaryDeleted][2]==0 or listPositions[i-boundaryDeleted][2]==pp.my_grid.bz):
        listPositions.pop(i-boundaryDeleted)
        boundaryDeleted+=1
#%% Ramdomly choose some positions where to place protein particles and create protein particles at these positions.
listInitialProteinParticles=[[],[]]
for i in range(2,pa.numProteinType):
    positionsRemaining=listPositions
    for j in range(i):
        positionsRemaining=list(set(positionsRemaining).difference(set(listInitialProteinParticles[j])))
    listInitialProteinParticles.append(random.sample(positionsRemaining,pa.N[i]))
    for position in listInitialProteinParticles[i]:
        pp.my_grid.grid[(position[0],position[1],position[2])]=pp.Protein[pa.listProteinNames[i]]
        std=math.sqrt(pa.k_B*pa.T/pa.m[i])
        vx,vy,vz=[random.normalvariate(0,std),random.normalvariate(0,std),random.normalvariate(0,std)]
        pp.listProteinParticles.append(pp.Particle(*position,vx,vy,vz,pp.Protein[pa.listProteinNames[i]]))
#%% Set input method.
inputFromFile=True
if inputFromFile:
    commandFile=open('command.txt')#Use the commands from 'command.txt' to direct the program running.
    figureNumber=0
    projectionNumber=0
    fillZero=5
    suffixFormat='.png'
    fps=30
    videoName='video.avi'
    outputPath='OutputFigures/'
    while os.path.exists(outputPath):
        outputPath=input('\nOutput Path \''+outputPath+'\' already exists!\nPlease input another path name: ')
        outputPath.replace('\\','')
        outputPath.replace('/','')
        outputPath+='/'
    os.makedirs(outputPath)
#%% Start simulation.
start_time=0
log=[]#To record the location, velocity, surrounding type, probability distribution, random number generated and direction determined of each protein particle at each time.
columns=['Particle ID','Time','x[0]','x[1]','x[2]','v[0]','v[1]','v[2]','t[0]','t[1]','t[2]','t[3]','t[4]','t[5]','t[6]','p[0]','p[1]','p[2]','p[3]','p[4]','p[5]','p[6]','Random Number Generated','Direction Determined']
while True:
    if inputFromFile:
        command=commandFile.readline()
        if command=='':
            command='q'
        command=command.replace('\n','')
    else:
        command=input('''
               Type a positive integer as the iteration time.
               Type "f" to start the FRAP experiment.
               Type "p" to print the current image of the system.
               Type "j" to print the projection image of the system.
               Type "l" to add light induction.
               Type "w" to withdraw light induction.
               Type "q" to quit the program.
        ''')
    if command in ['f','F']:
        pass#To be completed later.
    elif command in ['p','P']:
        if inputFromFile:
            figureNumber+=1
            savePath=outputPath+'fig'+str(figureNumber).zfill(fillZero)+suffixFormat
            vis.do_plot(pp.my_grid.grid,*pa.b,True,savePath)
            print('\n\''+savePath+'\' saved!\n')
        else:
            vis.do_plot(pp.my_grid.grid,*pa.b)
    elif command in ['j','J']:
        outputArray=proj.projection(pp.listProteinParticles,pp.my_grid,pa.projUnitSize,'z','max')
        if inputFromFile:
            projectionNumber+=1
            savePath=outputPath+'proj'+str(projectionNumber).zfill(fillZero)+suffixFormat
            proj.visualize(outputArray,vmin=0,vmax=4,saveFig=True,savePath=savePath)
            print('\n\''+savePath+'\' saved!\n')
        else:
            proj.visualize(outputArray)
    elif command in ['l','L']:
        pass#To be completed later.
    elif command in ['w','W']:
        pass#To be completed later.
    elif command in ['q','Q']:
        if inputFromFile:
            commandFile.close()
            if figureNumber>1:
                video_dir=outputPath+'fig_'+videoName#Start to creat a video file to show the process of protein movement.
                img=Image.open(outputPath+'fig'+'1'.zfill(fillZero)+suffixFormat)
                img_size=img.size
                fourcc=cv2.VideoWriter_fourcc('M','J','P','G')
                videoWriter=cv2.VideoWriter(video_dir,fourcc,fps,img_size)
                for i in range(figureNumber):
                    im_name=os.path.join(outputPath,'fig'+str(i+1).zfill(fillZero)+suffixFormat)
                    frame=cv2.imread(im_name)
                    videoWriter.write(frame)
                print('Video \''+'fig_'+videoName+'\' created!\n')
                videoWriter.release()#Video created.
            if projectionNumber>1:
                video_dir=outputPath+'proj'+videoName#Start to creat a video file to show the process of protein movement.
                img=Image.open(outputPath+'proj'+'1'.zfill(fillZero)+suffixFormat)
                img_size=img.size
                fourcc=cv2.VideoWriter_fourcc('M','J','P','G')
                videoWriter=cv2.VideoWriter(video_dir,fourcc,fps,img_size)
                for i in range(figureNumber):
                    im_name=os.path.join(outputPath,'proj'+str(i+1).zfill(fillZero)+suffixFormat)
                    frame=cv2.imread(im_name)
                    videoWriter.write(frame)
                print('Video \''+'proj'+videoName+'\' created!\n')
                videoWriter.release()#Video created.
            logPath=outputPath+'Log.xlsx'
            if input('The run has ended. Type \'c\' to continue the run, or anything else to exit the run: ') in ['c','C']:
                commandFile=open('command.txt')
                continue
        else:
            logPath=input('''Please enter the name for log file. Enter "skip" to skip: ''')
        if logPath.lower()=='skip':
            pass
        else:
            logPath=logPath.strip('.xlsx')+'.xlsx'
            dfLog=pd.DataFrame(log,columns=columns)
            dfLog.to_excel(logPath)#Record the log to a Microsoft Excel file.
        if inputFromFile:
            copyfile(pa.parameterFile,outputPath+pa.parameterFile)#Copy the parameter Excel file to output path.
            pa.record_parameters(outputPath+pa.parameterFile)#Record other parameters to the parameter Excel file.
            commentFile=docx.Document()
            comments=input('The run has ended. Write your comments here: ')
            commentFile.add_paragraph(comments)
            commentFile.add_paragraph('The following content is generated based on automatic translation:')
            commentFile.add_paragraph(Translator(to_lang='en',from_lang='zh').translate(comments))
            commentFile.save(outputPath+'Comments.docx')#Save comments to a Microsoft Word document.
        break
    else:
        validInput=True
        try:
            command=int(command)
            if (command<=0):
                validInput=False
        except ValueError:
            validInput=False
        if validInput:
            end_time=start_time+command
            while start_time<end_time:
                for i in [i for i in range(2,pa.numProteinType) if pa.N[i]>0]:
                    avgVelocity=[0,0,0]
                    for particle in [particle for particle in pp.listProteinParticles if particle.t[0]==pp.Protein[pa.listProteinNames[i]]]:
                        directionDetermined=particle.direction_determine()
                        log.append([id(particle),start_time,*tuple(particle.x),*tuple(particle.v),*tuple([str(particle.t[i]).strip('Protein.') for i in range(7)]),*tuple(particle.p),directionDetermined[0],directionDetermined[1]])
                        particle.move(directionDetermined[1])
                        avgVelocity=[avgVelocity[j]+particle.v[j] for j in range(3)]
                    avgVelocity=[avgVelocity[j]/pa.N[i] for j in range(3)]
                    varVelocity=0
                    for particle in [particle for particle in pp.listProteinParticles if particle.t[0]==pp.Protein[pa.listProteinNames[i]]]:
                        varVelocity+=sum([(particle.v[j]-avgVelocity[j])**2 for j in range(3)])
                    varVelocity/=pa.N[i]-1
                    stdCompensationVelocity=math.sqrt(max(0,pa.k_B*pa.T/pa.m[i]-varVelocity/3))
                    for particle in [particle for particle in pp.listProteinParticles if particle.t[0]==pp.Protein[pa.listProteinNames[i]]]:
                        particle.v=[particle.v[j]+random.normalvariate(0,stdCompensationVelocity) for j in range(3)]
                start_time+=1
        else:
            print('''
                  Invalid input!
            ''')